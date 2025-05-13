from django.conf import settings
from django.http import HttpResponse
from rest_framework import status, viewsets
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from rest_framework.exceptions import NotFound
import google.generativeai as genai
import io
import base64
from PIL import Image
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
import logging
import traceback
import json
from django.core.mail import EmailMessage
from accounts.utils import send_medical_report_email
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
from twilio.rest import Client
import os
import time
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RPImage, Table, TableStyle, ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER, TA_LEFT
import datetime

logger = logging.getLogger(__name__)

from .models import Report
from .serializers import ReportSerializer, ReportGenerateSerializer, ReportExportSerializer, ReportEmailSerializer


class ReportViewSet(viewsets.ReadOnlyModelViewSet):
    """ViewSet for viewing reports"""
    serializer_class = ReportSerializer
    permission_classes = [IsAuthenticated]
    
    def get_queryset(self):
        return Report.objects.filter(user=self.request.user)


class GenerateReportView(APIView):
    """Generate a medical report from an image"""
    permission_classes = [IsAuthenticated]
    
    def post(self, request):
        try:
            serializer = ReportGenerateSerializer(data=request.data)
            if not serializer.is_valid():
                logger.error("Serializer validation error: %s", serializer.errors)
                return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

            image_file = serializer.validated_data['image']
            language_code = serializer.validated_data.get('language', 'en')
            
            # Debug logging
            logger.info("API Key present: %s", bool(settings.GOOGLE_AI_API_KEY))
            logger.info("API Key length: %d", len(settings.GOOGLE_AI_API_KEY) if settings.GOOGLE_AI_API_KEY else 0)
            logger.info("Image file type: %s", type(image_file))
            logger.info("Image file name: %s", getattr(image_file, 'name', 'No name'))
            
            try:
                # Configure the Gemini AI model
                genai.configure(api_key=settings.GOOGLE_AI_API_KEY)
                
                # Get the model - using the newer version
                model = genai.GenerativeModel('gemini-1.5-flash')
                
                # Process the image
                try:
                    image_pil = Image.open(image_file)
                    logger.info("Image opened successfully. Size: %s", image_pil.size)
                except Exception as e:
                    logger.error("Error opening image: %s", str(e))
                    return Response({
                        'error': f"Error processing image: {str(e)}"
                    }, status=status.HTTP_400_BAD_REQUEST)
                
                # Generate report using Gemini
                try:
                    logger.info("Sending request to Gemini API...")
                    prompt = f"""Analyze this medical image and generate a detailed report in {language_code}. 
                    Please structure your response exactly as follows:

                    DIAGNOSIS:
                    [Write the main diagnosis or findings here]

                    DETAILS:
                    [Write detailed analysis here]

                    ACCURACY:
                    [Write a number between 0 and 1 indicating confidence level]

                    RECOMMENDATIONS:
                    [Write each recommendation on a new line with a bullet point]
                    """
                    
                    response = model.generate_content([prompt, image_pil])
                    logger.info("Received response from Gemini API")
                    
                    # Parse the response into structured data
                    report_text = response.text
                    logger.info("Raw AI response: %s", report_text)
                    
                    # Initialize default values
                    diagnosis = "No diagnosis available"
                    details = "No detailed analysis available"
                    accuracy = 0.8  # Default value
                    recommendations = ["No specific recommendations available"]
                    
                    # Split the response into sections
                    sections = report_text.split('\n\n')
                    current_section = None
                    
                    for section in sections:
                        section = section.strip()
                        if not section:
                            continue
                            
                        if section.upper().startswith('DIAGNOSIS:'):
                            current_section = 'diagnosis'
                            diagnosis = section.split(':', 1)[1].strip() if ':' in section else section.strip()
                        elif section.upper().startswith('DETAILS:'):
                            current_section = 'details'
                            details = section.split(':', 1)[1].strip() if ':' in section else section.strip()
                        elif section.upper().startswith('ACCURACY:'):
                            current_section = 'accuracy'
                            try:
                                accuracy_text = section.split(':', 1)[1].strip() if ':' in section else section.strip()
                                accuracy = float(accuracy_text)
                                # Ensure accuracy is between 0 and 1
                                accuracy = max(0.0, min(1.0, accuracy))
                            except ValueError:
                                accuracy = 0.8
                        elif section.upper().startswith('RECOMMENDATIONS:'):
                            current_section = 'recommendations'
                            recs = section.split(':', 1)[1].strip() if ':' in section else section.strip()
                            recommendations = [r.strip().lstrip('•- ') for r in recs.split('\n') if r.strip()]
                            if not recommendations:
                                recommendations = ["No specific recommendations available"]
                    
                    logger.info("Parsed sections - Diagnosis: %s, Details: %s, Accuracy: %s, Recommendations: %s",
                              diagnosis, details, accuracy, recommendations)
                    
                except Exception as e:
                    logger.error("Error from Gemini API: %s", str(e))
                    logger.error("Full traceback: %s", traceback.format_exc())
                    return Response({
                        'error': f"Error generating report with AI: {str(e)}"
                    }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                
                # Create a new report in the database
                try:
                    report = Report.objects.create(
                        user=request.user,
                        image=image_file,
                        diagnosis=diagnosis,
                        details=details,
                        accuracy=accuracy,
                        recommendations=json.dumps(recommendations),
                        language=language_code
                    )
                    logger.info("Report created in database with ID: %s", report.id)
                except Exception as e:
                    logger.error("Error saving report to database: %s", str(e))
                    return Response({
                        'error': f"Error saving report: {str(e)}"
                    }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                
                # Generate professional PDF report
                try:
                    pdf_buffer = self.generate_professional_pdf(request.user, diagnosis, details, accuracy, recommendations)
                    logger.info("Professional PDF report generated successfully")
                except Exception as e:
                    logger.error("Error generating PDF: %s", str(e))
                    return Response({
                        'error': f"Error generating PDF: {str(e)}"
                    }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                
                # Generate professional Word document
                try:
                    word_buffer = self.generate_professional_word_doc(request.user, diagnosis, details, accuracy, recommendations)
                    logger.info("Professional Word document generated successfully")
                except Exception as e:
                    logger.error("Error generating Word document: %s", str(e))
                    return Response({
                        'error': f"Error generating Word document: {str(e)}"
                    }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                
                # Return the report data
                return Response({
                    'report_id': report.id,
                    'diagnosis': diagnosis,
                    'details': details,
                    'accuracy': accuracy,
                    'recommendations': recommendations,
                    'pdf': base64.b64encode(pdf_buffer.getvalue()).decode('utf-8'),
                    'docx': base64.b64encode(word_buffer.getvalue()).decode('utf-8')
                }, status=status.HTTP_201_CREATED)
                
            except Exception as e:
                logger.error("Unexpected error: %s", str(e))
                logger.error("Full traceback: %s", traceback.format_exc())
                return Response({
                    'error': f"Unexpected error: {str(e)}"
                }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
                
        except Exception as e:
            logger.error("Top-level error: %s", str(e))
            logger.error("Full traceback: %s", traceback.format_exc())
            return Response({
                'error': f"Server error: {str(e)}"
            }, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def generate_professional_pdf(self, user, diagnosis, details, accuracy, recommendations):
        """Generate a professionally formatted PDF medical report"""
        from reportlab.lib.pagesizes import letter
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RPImage, Table, TableStyle, ListFlowable, ListItem
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        import datetime
        
        # Create a buffer to store the PDF
        buffer = io.BytesIO()
        
        # Set up the document with proper margins
        doc = SimpleDocTemplate(
            buffer, 
            pagesize=letter,
            rightMargin=72, 
            leftMargin=72,
            topMargin=72, 
            bottomMargin=72
        )
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # Create custom styles
        title_style = ParagraphStyle(
            'Title',
            parent=styles['Heading1'],
            fontSize=18,
            alignment=TA_CENTER,
            spaceAfter=12
        )
        
        heading_style = ParagraphStyle(
            'Heading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceBefore=12,
            spaceAfter=6
        )
        
        normal_style = ParagraphStyle(
            'Normal',
            parent=styles['Normal'],
            fontSize=11,
            spaceBefore=6,
            spaceAfter=6
        )
        
        # Build the document content
        elements = []
        
        # Add logo or hospital name if available - using placeholder text
        elements.append(Paragraph("Medical Analysis Report", title_style))
        
        # Add date and patient info
        current_date = datetime.datetime.now().strftime("%B %d, %Y")
        elements.append(Paragraph(f"Date: {current_date}", normal_style))
        elements.append(Paragraph(f"Patient: {user.get_full_name() or user.username}", normal_style))
        elements.append(Spacer(1, 0.25*inch))
        
        # Add horizontal line
        elements.append(Paragraph("<hr width='100%'/>", normal_style))
        elements.append(Spacer(1, 0.25*inch))
        
        # Add diagnosis section
        elements.append(Paragraph("Diagnosis", heading_style))
        elements.append(Paragraph(diagnosis, normal_style))
        elements.append(Spacer(1, 0.1*inch))
        
        # Add details section
        elements.append(Paragraph("Details", heading_style))
        elements.append(Paragraph(details, normal_style))
        elements.append(Spacer(1, 0.1*inch))
        
        # Add accuracy section with visual indicator
        elements.append(Paragraph("Diagnostic Confidence", heading_style))
        
        # Create a visual accuracy indicator
        accuracy_percentage = f"{accuracy * 100:.1f}%"
        accuracy_data = [
            ["Confidence Level:", accuracy_percentage]
        ]
        
        accuracy_table = Table(accuracy_data, colWidths=[2*inch, 1*inch])
        accuracy_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('FONTNAME', (0, 0), (0, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
        ]))
        
        elements.append(accuracy_table)
        elements.append(Spacer(1, 0.2*inch))
        
        # Add recommendations section
        elements.append(Paragraph("Recommendations", heading_style))
        
        # Create bullet points for recommendations
        recommendation_items = []
        for rec in recommendations:
            recommendation_items.append(ListItem(Paragraph(rec, normal_style)))
            
        recommendation_list = ListFlowable(
            recommendation_items,
            bulletType='bullet',
            start=None
        )
        elements.append(recommendation_list)
        
        # Add footer
        elements.append(Spacer(1, 0.5*inch))
        footer_text = "This report was generated using AI analysis and should be reviewed by a medical professional."
        elements.append(Paragraph(footer_text, ParagraphStyle(
            'Footer',
            parent=styles['Italic'],
            fontSize=9,
            textColor=colors.gray
        )))
        
        # Build the PDF
        doc.build(elements)
        buffer.seek(0)
        return buffer

    def generate_professional_word_doc(self, user, diagnosis, details, accuracy, recommendations):
        """Generate a professionally formatted Word document medical report"""
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.enum.style import WD_STYLE_TYPE
        import datetime
        
        # Create a new Document
        doc = Document()
        
        # Set document properties
        doc.core_properties.title = "Medical Analysis Report"
        doc.core_properties.author = "AI Medical System"
        
        # Define styles
        styles = doc.styles
        
        # Title style
        title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Arial'
        title_font.size = Pt(18)
        title_font.bold = True
        title_paragraph_format = title_style.paragraph_format
        title_paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_paragraph_format.space_after = Pt(12)
        
        # Heading style
        heading_style = styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
        heading_font = heading_style.font
        heading_font.name = 'Arial'
        heading_font.size = Pt(14)
        heading_font.bold = True
        heading_paragraph_format = heading_style.paragraph_format
        heading_paragraph_format.space_before = Pt(12)
        heading_paragraph_format.space_after = Pt(6)
        
        # Normal text style
        normal_style = styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
        normal_font = normal_style.font
        normal_font.name = 'Arial'
        normal_font.size = Pt(11)
        normal_paragraph_format = normal_style.paragraph_format
        normal_paragraph_format.space_before = Pt(6)
        normal_paragraph_format.space_after = Pt(6)
        normal_paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        # Add the title
        title = doc.add_paragraph("Medical Analysis Report", style='CustomTitle')
        
        # Add date and patient info
        current_date = datetime.datetime.now().strftime("%B %d, %Y")
        doc.add_paragraph(f"Date: {current_date}", style='CustomNormal')
        doc.add_paragraph(f"Patient: {user.get_full_name() or user.username}", style='CustomNormal')
        
        # Add horizontal line
        doc.add_paragraph().add_run().add_break()
        border_paragraph = doc.add_paragraph()
        border_paragraph.paragraph_format.border_bottom = True
        
        # Add diagnosis section
        doc.add_paragraph("Diagnosis", style='CustomHeading')
        doc.add_paragraph(diagnosis, style='CustomNormal')
        
        # Add details section
        doc.add_paragraph("Details", style='CustomHeading')
        doc.add_paragraph(details, style='CustomNormal')
        
        # Add accuracy section
        doc.add_paragraph("Diagnostic Confidence", style='CustomHeading')
        accuracy_percentage = f"{accuracy * 100:.1f}%"
        accuracy_para = doc.add_paragraph(style='CustomNormal')
        accuracy_para.add_run("Confidence Level: ").bold = True
        accuracy_para.add_run(accuracy_percentage)
        
        # Add recommendations section
        doc.add_paragraph("Recommendations", style='CustomHeading')
        
        # Add bullet points for recommendations
        for rec in recommendations:
            recommendation_para = doc.add_paragraph(style='CustomNormal')
            recommendation_para.style = 'List Bullet'
            recommendation_para.add_run(rec)
        
        # Add footer
        doc.add_paragraph().add_run().add_break()
        footer = doc.add_paragraph("This report was generated using AI analysis and should be reviewed by a medical professional.", style='CustomNormal')
        footer_run = footer.runs[0]
        footer_run.font.size = Pt(9)
        footer_run.font.italic = True
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
class SendReportEmailView(APIView):
    """Send a medical report by email"""
    permission_classes = [IsAuthenticated]
    
    def post(self, request):
        try:
            serializer = ReportEmailSerializer(data=request.data)
            if not serializer.is_valid():
                return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)
            
            report_id = serializer.validated_data['report_id']
            patient_email = serializer.validated_data['patient_email']
            patient_name = serializer.validated_data['patient_name']
            format_type = serializer.validated_data['format']
            
            try:
                report = Report.objects.get(id=report_id, user=request.user)
            except Report.DoesNotExist:
                return Response({'error': 'Report not found'}, status=status.HTTP_404_NOT_FOUND)
            
            # Generate the report file
            if format_type == 'pdf':
                # Generate PDF
                buffer = io.BytesIO()
                p = canvas.Canvas(buffer, pagesize=letter)
                p.drawString(100, 750, "Medical Report")
                p.drawString(100, 700, f"Diagnosis: {report.diagnosis}")
                p.drawString(100, 650, f"Details: {report.details}")
                p.drawString(100, 600, f"Accuracy: {report.accuracy * 100:.1f}%")
                p.drawString(100, 550, "Recommendations:")
                y = 500
                for rec in json.loads(report.recommendations):
                    p.drawString(120, y, f"• {rec}")
                    y -= 20
                p.save()
                report_file = buffer.getvalue()
            else:  # docx
                # Generate Word document
                doc = Document()
                doc.add_heading('Medical Report', 0)
                doc.add_heading('Diagnosis', level=1)
                doc.add_paragraph(report.diagnosis)
                doc.add_heading('Details', level=1)
                doc.add_paragraph(report.details)
                doc.add_heading('Accuracy', level=1)
                doc.add_paragraph(f"{report.accuracy * 100:.1f}%")
                doc.add_heading('Recommendations', level=1)
                for rec in json.loads(report.recommendations):
                    doc.add_paragraph(f"• {rec}", style='List Bullet')
                word_buffer = io.BytesIO()
                doc.save(word_buffer)
                report_file = word_buffer.getvalue()
            
            # Send email with the report
            send_medical_report_email(
                user=request.user,
                patient_email=patient_email,
                patient_name=patient_name,
                report_file=report_file,
                format_type=format_type
            )
            
            return Response({'message': 'Report sent successfully'}, status=status.HTTP_200_OK)
            
        except Exception as e:
            logger.error(f"Error sending report email: {str(e)}")
            return Response({'error': f"Error sending email: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


@csrf_exempt
def send_whatsapp_message(request):
    if request.method == 'POST':
        try:
            # Log the incoming request
            logger.info("Received WhatsApp message request")
            
            data = json.loads(request.body)
            phone = data.get('phone')
            message_text = data.get('message')
            pdf_base64 = data.get('pdf')  # Base64 encoded PDF

            logger.info(f"Request data - Phone: {phone}, Has message: {bool(message_text)}, Has PDF: {bool(pdf_base64)}")

            if not phone:
                logger.error("Phone number missing in request")
                return JsonResponse({
                    'status': 'error',
                    'message': 'Phone number is required'
                }, status=400)

            if not message_text and not pdf_base64:
                logger.error("Both message and PDF are missing")
                return JsonResponse({
                    'status': 'error',
                    'message': 'Either message text or PDF is required'
                }, status=400)

            # Clean and validate phone number
            phone = phone.strip()
            if not phone.startswith('+'):
                phone = '+' + phone
            
            # Remove any spaces or special characters
            phone = ''.join(c for c in phone if c.isdigit() or c == '+')
            
            # Basic phone number validation
            if len(phone) < 10 or len(phone) > 15:
                logger.error(f"Invalid phone number format: {phone}")
                return JsonResponse({
                    'status': 'error',
                    'message': 'Invalid phone number format. Please provide a valid international number (e.g., +1234567890)'
                }, status=400)

            account_sid = os.environ.get('TWILIO_ACCOUNT_SID')
            auth_token = os.environ.get('TWILIO_AUTH_TOKEN')
            from_whatsapp_number = os.environ.get('TWILIO_WHATSAPP_NUMBER')

            # Check if trying to send to the same number
            if phone == from_whatsapp_number.replace('whatsapp:', ''):
                logger.error("Attempting to send message to the same number")
                return JsonResponse({
                    'status': 'error',
                    'message': 'Cannot send message to the same number. Please provide a different recipient number.'
                }, status=400)

            try:
                client = Client(account_sid, auth_token)
                
                # Log the attempt
                logger.info(f"Attempting to send WhatsApp message to {phone}")
                
                # First, send the text message
                text_message_params = {
                    'from_': from_whatsapp_number,
                    'to': f'whatsapp:{phone}',
                    'body': message_text
                }

                logger.info("Sending text message first")
                text_message = client.messages.create(**text_message_params)
                logger.info(f"Text message sent successfully. SID: {text_message.sid}")

                # If PDF is provided, try to send it as a separate message
                if pdf_base64:
                    try:
                        logger.info("Processing PDF file")
                        # Decode base64 PDF
                        pdf_data = base64.b64decode(pdf_base64)
                        
                        # Create temp directory if it doesn't exist
                        temp_dir = os.path.join(settings.MEDIA_ROOT, 'temp')
                        os.makedirs(temp_dir, exist_ok=True)
                        
                        # Save PDF temporarily
                        temp_pdf_path = os.path.join(temp_dir, f'whatsapp_pdf_{int(time.time())}.pdf')
                        logger.info(f"Saving PDF to: {temp_pdf_path}")
                        
                        with open(temp_pdf_path, 'wb') as f:
                            f.write(pdf_data)
                        
                        # Get the media URL - use the full domain
                        domain = request.build_absolute_uri('/').rstrip('/')
                        media_url = f"{domain}{settings.MEDIA_URL}temp/{os.path.basename(temp_pdf_path)}"
                        logger.info(f"PDF saved successfully. Media URL: {media_url}")

                        # Try to send PDF as a separate message
                        pdf_message_params = {
                            'from_': from_whatsapp_number,
                            'to': f'whatsapp:{phone}',
                            'body': 'Here is your medical report PDF:',
                            'media_url': [media_url]
                        }

                        logger.info("Attempting to send PDF")
                        pdf_message = client.messages.create(**pdf_message_params)
                        logger.info(f"PDF message sent successfully. SID: {pdf_message.sid}")

                        return JsonResponse({
                            'status': 'success',
                            'sid': text_message.sid,
                            'pdf_sid': pdf_message.sid,
                            'message': 'WhatsApp message and PDF sent successfully'
                        })

                    except Exception as pdf_error:
                        logger.error(f"Error sending PDF: {str(pdf_error)}")
                        # If PDF sending fails, we still return success for the text message
                        return JsonResponse({
                            'status': 'partial_success',
                            'sid': text_message.sid,
                            'message': 'Text message sent successfully, but PDF could not be sent. Please download the report from the web interface.'
                        })

                # If no PDF, just return success for text message
                return JsonResponse({
                    'status': 'success',
                    'sid': text_message.sid,
                    'message': 'WhatsApp message sent successfully'
                })

            except Exception as twilio_error:
                error_message = str(twilio_error)
                logger.error(f"Twilio API error: {error_message}")
                logger.error(f"Full traceback: {traceback.format_exc()}")
                
                # Handle specific Twilio error codes
                if '63031' in error_message:
                    return JsonResponse({
                        'status': 'error',
                        'message': '''Invalid WhatsApp number. To use this service:
1. The recipient must first join the Twilio WhatsApp sandbox
2. Send "join <two-words>" to the WhatsApp number: +212669387200
3. Wait for the confirmation message
4. Then try sending your message again'''
                    }, status=400)
                elif '21211' in error_message:
                    return JsonResponse({
                        'status': 'error',
                        'message': 'Invalid phone number format. Please provide a valid international number.'
                    }, status=400)
                else:
                    return JsonResponse({
                        'status': 'error',
                        'message': f'Twilio API error: {error_message}'
                    }, status=500)

        except json.JSONDecodeError as e:
            logger.error(f"JSON decode error: {str(e)}")
            return JsonResponse({
                'status': 'error',
                'message': 'Invalid JSON data'
            }, status=400)
        except Exception as e:
            logger.error(f"Unexpected error: {str(e)}")
            logger.error(f"Full traceback: {traceback.format_exc()}")
            return JsonResponse({
                'status': 'error',
                'message': f'Server error: {str(e)}'
            }, status=500)

    return JsonResponse({
        'status': 'error',
        'message': 'Only POST method allowed'
    }, status=405)